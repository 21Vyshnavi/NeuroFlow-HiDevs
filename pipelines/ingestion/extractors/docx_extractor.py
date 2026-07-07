import io
import logging
from typing import List
from pipelines.ingestion.extractors.pdf_extractor import ExtractedPage

logger = logging.getLogger(__name__)

async def extract_docx(file_bytes: bytes) -> List[ExtractedPage]:
    pages = []
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        heading_metadata = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                if para.style.name.startswith("Heading"):
                    heading_metadata.append({"style": para.style.name, "text": para.text})

        # Simple table cell extractor
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        pages.append(ExtractedPage(
            page_number=1,
            content="\n\n".join(full_text),
            content_type="text",
            metadata={"headings": heading_metadata}
        ))
    except Exception as e:
        logger.warning(f"DOCX Extractor failed: {e}")
        pages.append(ExtractedPage(
            page_number=1,
            content=file_bytes.decode('utf-8', errors='ignore'),
            content_type="text",
            metadata={"engine": "fallback"}
        ))
    return pages
